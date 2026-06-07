"""
Сборка и дополнение файла ВКР cursor.docx на рабочем столе.
Оформление — точь-в-точь по шаблону Диплом1.docx (клонирование абзацев).
"""
from __future__ import annotations

import argparse
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from vkr_content import (
    ANALOGUES_SECTION,
    ANNOTATION_EN,
    ANNOTATION_RU,
    LISTINGS,
    PRACTICE_CONCLUSION_BLOCK,
    RELEVANCE_SECTION,
    SOURCES,
    TOOLS_SECTION,
)
from vkr_practice_sections import PRACTICE_SECTIONS
from vkr_template_format import TemplateFormat

TEMPLATE = Path(r"c:\Users\Klim\Documents\Диплом образцы\Диплом1.docx")
OUTPUT = Path(r"c:\Users\Klim\Desktop\ВКР cursor.docx")
FONT_NAME = "Liberation Serif"
CODE_FONT = "Consolas"

_FMT = TemplateFormat()


def clear_body_keep_section(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def ensure_document(*, rebuild: bool = False) -> Document:
    if rebuild and OUTPUT.exists():
        try:
            OUTPUT.unlink()
        except OSError:
            doc = Document(str(OUTPUT))
            clear_body_keep_section(doc)
            return doc
    if OUTPUT.exists():
        return Document(str(OUTPUT))
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Шаблон не найден: {TEMPLATE}")
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))
    clear_body_keep_section(doc)
    return doc


def _set_run_font(run, *, size_pt: float, font: str = FONT_NAME, bold: bool = False) -> None:
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.bold = bold
    if font == FONT_NAME:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            from docx.oxml import OxmlElement

            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), FONT_NAME)
        rfonts.set(qn("w:hAnsi"), FONT_NAME)
        rfonts.set(qn("w:cs"), FONT_NAME)


def add_code_block(doc: Document, code: str) -> None:
    """Листинги — отдельный стиль, в шаблоне для кода используется моноширинный шрифт."""
    for line in code.splitlines():
        p = doc.add_paragraph(style="Без отступа текста")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line if line else " ")
        _set_run_font(run, size_pt=10, font=CODE_FONT)


def add_annotation_ru(doc: Document) -> None:
    _FMT.add_annotation_heading_ru(doc)
    _FMT.add_body(doc, ANNOTATION_RU["meta"])
    _FMT.add_body(doc, ANNOTATION_RU["object"])
    _FMT.add_body(doc, ANNOTATION_RU["goal"])
    _FMT.add_body(doc, ANNOTATION_RU["tools"])
    _FMT.add_body(doc, "Задачи выпускной квалификационной работы включают:")
    for task in ANNOTATION_RU["tasks"]:
        _FMT.add_task_ru(doc, task)
    _FMT.add_body(doc, ANNOTATION_RU["result"])
    _FMT.add_body_ru_area_end(doc, ANNOTATION_RU["area"])


def add_annotation_en(doc: Document) -> None:
    _FMT.add_annotation_heading_en(doc)
    _FMT.add_body(doc, ANNOTATION_EN["meta"])
    _FMT.add_body(doc, ANNOTATION_EN["object"])
    _FMT.add_body(doc, ANNOTATION_EN["goal"])
    _FMT.add_body(doc, ANNOTATION_EN["tools"])
    _FMT.add_body(doc, "The tasks of the final qualification work include:")
    for task in ANNOTATION_EN["tasks"]:
        _FMT.add_task_en(doc, task)
    _FMT.add_body(doc, ANNOTATION_EN["result"])
    _FMT.add_body_en_area_end(doc, ANNOTATION_EN["area"])


def render_section(doc: Document, section: dict) -> None:
    level = section.get("level", 2)
    if level == 1:
        _FMT.add_chapter(doc, section["title"])
    elif level == 2:
        _FMT.add_subsection(doc, section["title"])
    else:
        _FMT.add_subsection(doc, section["title"])

    for para in section.get("paragraphs", []):
        _FMT.add_section_body(doc, para)

    for bullet in section.get("bullets", []):
        _FMT.add_task_ru(doc, bullet)

    if section.get("use_practice_block"):
        for block in PRACTICE_CONCLUSION_BLOCK.split("\n\n"):
            block = block.strip()
            if block:
                _FMT.add_section_body(doc, block)

    for para in section.get("extra_paragraphs", []):
        _FMT.add_section_body(doc, para)

    for sub in section.get("subsections", []):
        _FMT.add_subsection(doc, sub["title"])
        for para in sub.get("paragraphs", []):
            _FMT.add_section_body(doc, para)
        for bullet in sub.get("bullets", []):
            _FMT.add_task_ru(doc, bullet)

    if section.get("sources"):
        for i, source in enumerate(SOURCES, start=1):
            _FMT.add_body(doc, f"{i}. {source}")

    if section.get("listings"):
        for title, code, caption in LISTINGS:
            _FMT.add_body(doc, title)
            add_code_block(doc, code)
            _FMT.add_body(doc, caption)


def _preserve_tail_elements(doc: Document) -> list:
    tail: list = []
    tail_started = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if RELEVANCE_SECTION["chapter"] in text:
            tail_started = True
        if tail_started:
            tail.append(deepcopy(paragraph._p))
    return tail


def rebuild_annotation(*, preserve_tail: bool = True) -> Path:
    """Пересобрать аннотацию (2 страницы: RU + EN с разрывом), сохранить разделы после неё."""
    doc = ensure_document(rebuild=False)
    tail = _preserve_tail_elements(doc) if preserve_tail else []
    clear_body_keep_section(doc)
    add_annotation_ru(doc)
    add_annotation_en(doc)
    for element in tail:
        doc.element.body.append(element)
    return save(doc)


def build_full_document(*, rebuild: bool = True) -> Path:
    doc = ensure_document(rebuild=rebuild)
    add_annotation_ru(doc)
    add_annotation_en(doc)
    return save(doc)


def _split_before_relevance(doc: Document) -> list:
    before: list = []
    for paragraph in doc.paragraphs:
        if RELEVANCE_SECTION["chapter"] in paragraph.text:
            break
        before.append(deepcopy(paragraph._p))
    return before


def _split_before_subsection(doc: Document, prefix: str) -> list:
    before: list = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(prefix):
            break
        before.append(deepcopy(paragraph._p))
    return before


def _add_comparison_table(doc: Document, table_data: dict) -> None:
    columns = table_data["columns"]
    rows_data = table_data["rows"]
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(columns) + 1)
    table.style = "Table Grid"

    header = table.rows[0].cells
    header[0].text = "Критерий"
    for index, name in enumerate(columns):
        header[index + 1].text = name

    for row_index, (criterion, marks) in enumerate(rows_data):
        row = table.rows[row_index + 1].cells
        row[0].text = criterion
        for col_index, mark in enumerate(marks):
            row[col_index + 1].text = mark


def _render_relevance_section(doc: Document) -> None:
    _FMT.add_chapter(doc, RELEVANCE_SECTION["chapter"])
    _FMT.add_empty(doc)
    _FMT.add_subsection(doc, RELEVANCE_SECTION["subsection"])
    for para in RELEVANCE_SECTION["paragraphs"]:
        _FMT.add_section_body(doc, para)


def replace_relevance_section() -> Path:
    """Заменить раздел 1.1 Актуальность, сохранив аннотацию."""
    doc = ensure_document(rebuild=False)
    before = _split_before_relevance(doc)
    clear_body_keep_section(doc)
    for element in before:
        doc.element.body.append(element)
    _render_relevance_section(doc)
    return save(doc)


def append_relevance_section() -> Path:
    doc = ensure_document(rebuild=False)
    if any(RELEVANCE_SECTION["chapter"] in p.text for p in doc.paragraphs):
        return replace_relevance_section()
    _render_relevance_section(doc)
    return save(doc)


def _render_analogues_section(doc: Document) -> None:
    _FMT.add_subsection(doc, ANALOGUES_SECTION["subsection"])
    for paragraph in ANALOGUES_SECTION["intro"]:
        _FMT.add_section_body(doc, paragraph)

    for index, analogue in enumerate(ANALOGUES_SECTION["analogues"], start=1):
        _FMT.add_analogue_title(doc, index, analogue["name"])
        for paragraph in analogue["description"]:
            _FMT.add_section_body(doc, paragraph)
        _FMT.add_section_body(doc, analogue["pros"])
        _FMT.add_section_body(doc, analogue["cons"])

    for paragraph in ANALOGUES_SECTION["summary"]:
        _FMT.add_section_body(doc, paragraph)

    _FMT.add_table_caption(doc, ANALOGUES_SECTION["table_caption"])
    _add_comparison_table(doc, ANALOGUES_SECTION["table"])
    _FMT.add_section_body(doc, ANALOGUES_SECTION["requirements_intro"])
    for requirement in ANALOGUES_SECTION["requirements"]:
        _FMT.add_section_body(doc, requirement)


def append_analogues_section() -> Path:
    doc = ensure_document(rebuild=False)
    before = _split_before_subsection(doc, "1.2")
    clear_body_keep_section(doc)
    for element in before:
        doc.element.body.append(element)
    _render_analogues_section(doc)
    return save(doc)


def _render_tools_section(doc: Document) -> None:
    _FMT.add_subsection(doc, TOOLS_SECTION["subsection"])
    for paragraph in TOOLS_SECTION["intro"]:
        _FMT.add_section_body(doc, paragraph)

    for technology in TOOLS_SECTION["technologies"]:
        _FMT.add_section_body(doc, technology["name"])
        _FMT.add_section_body(doc, "Достоинства:")
        for item in technology["pros"]:
            _FMT.add_section_body(doc, item)
        _FMT.add_section_body(doc, "Недостатки:")
        for item in technology["cons"]:
            _FMT.add_section_body(doc, item)

    for paragraph in TOOLS_SECTION["conclusion"]:
        _FMT.add_section_body(doc, paragraph)

    _FMT.add_section_body(doc, TOOLS_SECTION["stack_intro"])
    for item in TOOLS_SECTION["stack"]:
        _FMT.add_section_body(doc, item)

    for paragraph in TOOLS_SECTION.get("components_intro", []):
        _FMT.add_section_body(doc, paragraph)

    for component in TOOLS_SECTION.get("components", []):
        _FMT.add_section_body(doc, component["name"])
        for paragraph in component["paragraphs"]:
            _FMT.add_section_body(doc, paragraph)


def append_tools_section() -> Path:
    doc = ensure_document(rebuild=False)
    before = _split_before_subsection(doc, "1.3")
    clear_body_keep_section(doc)
    for element in before:
        doc.element.body.append(element)
    _render_tools_section(doc)
    return save(doc)


def append_practice_section(section: dict) -> Path:
    doc = ensure_document(rebuild=False)
    render_section(doc, section)
    return save(doc)


def append_text(heading: str | None, paragraphs: list[str], *, level: int = 2) -> None:
    doc = ensure_document(rebuild=False)
    if heading:
        if level == 1:
            _FMT.add_chapter(doc, heading)
        else:
            _FMT.add_subsection(doc, heading)
    for para in paragraphs:
        _FMT.add_section_body(doc, para)
    return save(doc)


def save(doc: Document) -> Path:
    candidates = [
        OUTPUT,
        OUTPUT.with_name(f"{OUTPUT.stem} — обновлён{OUTPUT.suffix}"),
        OUTPUT.with_name(f"{OUTPUT.stem} — полный{OUTPUT.suffix}"),
        OUTPUT.with_name(f"{OUTPUT.stem} — аннотация{OUTPUT.suffix}"),
        OUTPUT.with_name(f"{OUTPUT.stem} — актуальность{OUTPUT.suffix}"),
        OUTPUT.with_name(f"{OUTPUT.stem} — аналоги{OUTPUT.suffix}"),
        OUTPUT.with_name(f"{OUTPUT.stem} — средства{OUTPUT.suffix}"),
    ]
    for path in candidates:
        try:
            doc.save(str(path))
            if path != OUTPUT:
                print(
                    f"Основной файл занят — сохранено в: {path}\n"
                    "Закройте ВКР cursor.docx в Word для записи в основной файл.",
                    file=sys.stderr,
                )
            return path
        except PermissionError:
            continue
    raise PermissionError(f"Не удалось сохранить документ: {OUTPUT}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Сборка ВКР cursor.docx")
    parser.add_argument("--rebuild", action="store_true", help="Пересоздать документ из шаблона")
    parser.add_argument(
        "--annotation-only",
        action="store_true",
        help="Пересобрать только аннотацию, сохранив последующие разделы",
    )
    parser.add_argument(
        "--relevance",
        action="store_true",
        help="Обновить раздел 1.1 Актуальность",
    )
    parser.add_argument(
        "--analogues",
        action="store_true",
        help="Добавить или обновить раздел 1.2 Обзор аналогов",
    )
    parser.add_argument(
        "--tools",
        action="store_true",
        help="Добавить или обновить раздел 1.3 Выбор программных средств",
    )
    args = parser.parse_args()
    if args.tools:
        path = append_tools_section()
    elif args.analogues:
        path = append_analogues_section()
    elif args.relevance:
        path = replace_relevance_section()
    elif args.annotation_only:
        path = rebuild_annotation()
    else:
        path = build_full_document(rebuild=args.rebuild or not OUTPUT.exists())
    print(f"Готово: {path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
