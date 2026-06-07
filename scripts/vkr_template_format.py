"""
Форматирование по образцу Диплом1.docx — клонирование абзацев шаблона.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATE = Path(r"c:\Users\Klim\Documents\Диплом образцы\Диплом1.docx")

# Индексы абзацев-образцов в шаблоне
IDX_ANNOTATION_HEADING_RU = 102
IDX_BODY = 103
IDX_TASK_RU = 108
IDX_ANNOTATION_HEADING_EN = 116
IDX_TASK_EN = 122
IDX_BODY_RU_AREA_END = 115
IDX_BODY_EN_AREA_END = 129
IDX_CHAPTER = 152
IDX_EMPTY_AFTER_CHAPTER = 153
IDX_SUBSECTION = 154
IDX_BODY_AFTER_SUBSECTION = 155
IDX_ANALOGUE_TITLE = 164


class TemplateFormat:
    def __init__(self, template_path: Path = TEMPLATE) -> None:
        self._template = Document(str(template_path))

    def add_clone(self, doc: Document, index: int, text: str, *, keep_page_break: bool = False) -> None:
        src = self._template.paragraphs[index]._p
        new_p = deepcopy(src)
        if keep_page_break:
            self._replace_text_keep_page_break(new_p, text, src)
        else:
            self._replace_text(new_p, text, src)
        doc.element.body.append(new_p)

    def _replace_text(self, p, text: str, src_p) -> None:
        for child in list(p):
            if child.tag == qn("w:pPr"):
                continue
            p.remove(child)

        r_pr = self._copy_r_pr(src_p)
        r = OxmlElement("w:r")
        if r_pr is not None:
            r.append(r_pr)
        t = OxmlElement("w:t")
        if text.startswith(" ") or text.endswith(" "):
            t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)

    def _replace_text_keep_page_break(self, p, text: str, src_p) -> None:
        page_break_runs = []
        for child in src_p:
            if child.tag != qn("w:r"):
                continue
            br = child.find(qn("w:br"))
            if br is not None and br.get(qn("w:type")) == "page":
                page_break_runs.append(deepcopy(child))

        self._replace_text(p, text, src_p)
        for br_run in page_break_runs:
            p.append(br_run)

    def _copy_r_pr(self, src_p):
        for child in src_p:
            if child.tag == qn("w:r"):
                r_pr = child.find(qn("w:rPr"))
                if r_pr is not None:
                    return deepcopy(r_pr)
        p_pr = src_p.find(qn("w:pPr"))
        if p_pr is not None:
            r_pr = p_pr.find(qn("w:rPr"))
            if r_pr is not None:
                return deepcopy(r_pr)
        return None

    def add_annotation_heading_ru(self, doc: Document, text: str = "АННОТАЦИЯ") -> None:
        self.add_clone(doc, IDX_ANNOTATION_HEADING_RU, text)

    def add_annotation_heading_en(self, doc: Document, text: str = "THE ANNOTATION") -> None:
        self.add_clone(doc, IDX_ANNOTATION_HEADING_EN, text)

    def add_body(self, doc: Document, text: str) -> None:
        self.add_clone(doc, IDX_BODY, text)

    def add_body_ru_area_end(self, doc: Document, text: str) -> None:
        """Последний абзац русской аннотации — с разрывом страницы перед THE ANNOTATION."""
        self.add_clone(doc, IDX_BODY_RU_AREA_END, text, keep_page_break=True)

    def add_body_en_area_end(self, doc: Document, text: str) -> None:
        self.add_clone(doc, IDX_BODY_EN_AREA_END, text)

    def add_task_ru(self, doc: Document, text: str) -> None:
        self.add_clone(doc, IDX_TASK_RU, text)

    def add_task_en(self, doc: Document, text: str) -> None:
        bullet = "\uf02d"
        self.add_clone(doc, IDX_TASK_EN, f"{bullet}\t{text}")

    def add_chapter(self, doc: Document, text: str) -> None:
        self.add_clone(doc, IDX_CHAPTER, text)

    def add_empty(self, doc: Document) -> None:
        self.add_clone(doc, IDX_EMPTY_AFTER_CHAPTER, "")

    def add_subsection(self, doc: Document, text: str) -> None:
        self.add_clone(doc, IDX_SUBSECTION, text)

    def add_section_body(self, doc: Document, text: str) -> None:
        self.add_clone(doc, IDX_BODY_AFTER_SUBSECTION, text)

    def add_analogue_title(self, doc: Document, number: int, name: str) -> None:
        self.add_clone(doc, IDX_ANALOGUE_TITLE, f"{number}.\t{name}.")

    def add_table_caption(self, doc: Document, text: str) -> None:
        self.add_clone(doc, IDX_BODY, text)
