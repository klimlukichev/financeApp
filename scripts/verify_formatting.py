"""Сравнение форматирования output vs template для блока аннотации."""
import re
import zipfile
from pathlib import Path

from docx import Document

TEMPLATE = Path(r"c:\Users\Klim\Documents\Диплом образцы\Диплом1.docx")
OUTPUT = Path(r"c:\Users\Klim\Desktop\ВКР cursor.docx")


def ppr_signature(p) -> str:
    xml = p._p.xml
    ppr = re.search(r"<w:pPr>.*?</w:pPr>", xml, re.S)
    return re.sub(r"<w:t[^>]*>.*?</w:t>", "", ppr.group(0) if ppr else "")


def main() -> None:
    tdoc = Document(str(TEMPLATE))
    odoc = Document(str(OUTPUT))

    pairs = [
        (102, 0, "heading RU"),
        (103, 1, "meta"),
        (108, 6, "task RU"),
        (116, 14, "heading EN"),
        (122, 20, "task EN"),
    ]

    ok = True
    for t_idx, o_idx, label in pairs:
        if o_idx >= len(odoc.paragraphs):
            print(f"MISSING {label}: output para {o_idx}")
            ok = False
            continue
        ts = ppr_signature(tdoc.paragraphs[t_idx])
        os_ = ppr_signature(odoc.paragraphs[o_idx])
        match = ts == os_
        ok = ok and match
        print(f"{label}: {'OK' if match else 'DIFF'}")
        if not match:
            print("  template:", ts[:200])
            print("  output:  ", os_[:200])

    print(f"\nOutput paragraphs: {len(odoc.paragraphs)} (expected 28)")
    print("ALL OK" if ok and len(odoc.paragraphs) == 28 else "CHECK NEEDED")


if __name__ == "__main__":
    main()
